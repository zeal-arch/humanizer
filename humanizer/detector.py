# humanizer/detector.py
# Rule-based and Transformer-based AI probability scorer
# Uses DeBERTa-v3 fine-tuned model for deep classification and fallbacks to heuristics

import re
import math
import os
import threading
from .phrases import AI_PHRASES, INTENSIFIERS

try:
    from transformers import PreTrainedModel, AutoConfig, AutoModel, AutoTokenizer
    import torch
    import torch.nn as nn
    HAS_TORCH_DEPS = True
except ImportError:
    PreTrainedModel = object  # dummy base class for syntax validation
    nn = object
    HAS_TORCH_DEPS = False

if HAS_TORCH_DEPS:
    class DesklibAIDetectionModel(PreTrainedModel):
        config_class = AutoConfig

        def __init__(self, config):
            super().__init__(config)
            self.model = AutoModel.from_config(config)
            self.classifier = nn.Linear(config.hidden_size, 1)
            self.post_init()

        def forward(self, input_ids, attention_mask=None, labels=None):
            outputs = self.model(input_ids, attention_mask=attention_mask)
            last_hidden_state = outputs[0]
            input_mask_expanded = attention_mask.unsqueeze(-1).expand(last_hidden_state.size()).float()
            sum_embeddings = torch.sum(last_hidden_state * input_mask_expanded, dim=1)
            sum_mask = torch.clamp(input_mask_expanded.sum(dim=1), min=1e-9)
            pooled_output = sum_embeddings / sum_mask

            logits = self.classifier(pooled_output)
            loss = None
            if labels is not None:
                loss_fct = nn.BCEWithLogitsLoss()
                loss = loss_fct(logits.view(-1), labels.float())

            output = {"logits": logits}
            if loss is not None:
                output["loss"] = loss
            return output
else:
    class DesklibAIDetectionModel:
        pass


# ─────────────────────────────────────────────────────────────────────────────
# SIGNAL 1 — AI Phrase Density
# Count how many known AI phrases exist per 100 words
# ─────────────────────────────────────────────────────────────────────────────

def _ai_phrase_density(text: str) -> float:
    """
    Returns a 0.0–1.0 score based on how many AI phrases per 100 words.
    Saturates at ~12 hits per 100 words = 1.0
    """
    word_count = max(len(text.split()), 1)
    hits = 0
    combined = {**AI_PHRASES, **INTENSIFIERS}
    lower_text = text.lower()
    for phrase in combined:
        count = lower_text.count(phrase.lower())
        hits += count

    # Normalize: 12 hits per 100 words → score of 1.0
    density = (hits / word_count) * 100
    return min(density / 12.0, 1.0)


# ─────────────────────────────────────────────────────────────────────────────
# SIGNAL 2 — Burstiness (sentence length variance)
# Humans have high variance in sentence length. AI flatlines.
# ─────────────────────────────────────────────────────────────────────────────

def _tokenize_sentences(text: str) -> list[str]:
    parts = re.split(r'(?<=[.!?])\s+(?=[A-Z])', text)
    return [p.strip() for p in parts if p.strip() and len(p.split()) > 2]


def _burstiness_score(text: str) -> float:
    """
    Returns 0.0–1.0 where 1.0 = very AI-like (low burstiness).
    Low std dev in sentence length → higher AI score.
    """
    sentences = _tokenize_sentences(text)
    if len(sentences) < 4:
        return 0.5  # not enough data

    lengths = [len(s.split()) for s in sentences]
    mean = sum(lengths) / len(lengths)
    variance = sum((l - mean) ** 2 for l in lengths) / len(lengths)
    std_dev = math.sqrt(variance)

    # Human text typically has std_dev > 12 words
    # AI text typically has std_dev < 7 words
    # Map: std_dev=0 → score=1.0, std_dev=15 → score=0.0
    ai_score = max(0.0, min(1.0, 1.0 - (std_dev / 15.0)))
    return ai_score


# ─────────────────────────────────────────────────────────────────────────────
# SIGNAL 3 — Paragraph Opener Uniformity
# AI starts many paragraphs with "The", "This", "In", "It"
# ─────────────────────────────────────────────────────────────────────────────

AI_OPENERS = {
    "the ", "this ", "these ", "those ", "in ", "it ", "by ", "with ",
    "as ", "one ", "another ", "such ", "each ", "every "
}

def _opener_uniformity(text: str) -> float:
    """
    Returns 0.0–1.0 where 1.0 = all paragraphs start with AI opener words.
    """
    paragraphs = [p.strip() for p in text.split('\n') if len(p.strip().split()) >= 5]
    if not paragraphs:
        return 0.5

    ai_starts = sum(
        1 for p in paragraphs
        if any(p.lower().startswith(op) for op in AI_OPENERS)
    )
    return ai_starts / max(len(paragraphs), 1)


# ─────────────────────────────────────────────────────────────────────────────
# SIGNAL 4 — Average sentence length
# AI tends to write sentences of 20–35 words, very consistently
# Humans write shorter AND longer sentences
# ─────────────────────────────────────────────────────────────────────────────

def _avg_sentence_length_score(text: str) -> float:
    """
    Returns 0.0–1.0. Sentences clustered between 20-35 words = AI-like.
    """
    sentences = _tokenize_sentences(text)
    if not sentences:
        return 0.5

    lengths = [len(s.split()) for s in sentences]
    avg = sum(lengths) / len(lengths)

    # AI sweet spot: 22–32 words avg → score near 1.0
    # Very short (<12) or very long (>40) → more human-like
    if 20 <= avg <= 35:
        return 0.85
    elif 15 <= avg < 20 or 35 < avg <= 40:
        return 0.55
    else:
        return 0.2


# ─────────────────────────────────────────────────────────────────────────────
# SIGNAL 5 — Intensifier / filler word density
# ─────────────────────────────────────────────────────────────────────────────

FILLER_WORDS = [
    "very", "quite", "extremely", "highly", "incredibly", "absolutely",
    "essentially", "fundamentally", "primarily", "significantly", "substantially",
    "undoubtedly", "undeniably", "seamlessly", "effectively", "efficiently",
    "comprehensive", "robust", "innovative", "leverage", "leveraged", "leveraging",
    "leverages", "utilize", "utilized", "utilizing", "utilizes", "ensure",
    "ensured", "ensuring", "ensures", "facilitate", "facilitated", "facilitating",
    "facilitates", "facilitation", "demonstrate", "demonstrated", "demonstrating",
    "demonstrates", "demonstration", "demonstrations", "implement", "implemented",
    "implementing", "implements", "implementation", "implementations", "paramount",
    "pivotal",
    # Modern AI cliches (GPT-4o, GPT-5, Claude Opus, Gemini) and their variations
    "delve", "delves", "delved", "delving", "tapestry", "tapestries", "myriad",
    "orchestrate", "orchestrates", "orchestrated", "orchestrating", "orchestration",
    "orchestrations", "catalyst", "catalysts", "foster", "fosters", "fostered",
    "fostering", "unleash", "unleashes", "unleashed", "unleashing", "empower",
    "empowers", "empowered", "empowering", "empowerment", "empowerments", "beacon",
    "beacons", "spearhead", "spearheads", "spearheaded", "spearheading", "meticulously",
    "harmonious", "synergize", "synergizes", "synergized", "synergizing", "synergy",
    "synergies", "frictionless", "actionable", "catalyze", "catalyzes", "catalyzed",
    "catalyzing", "democratize", "democratizes", "democratized", "democratizing",
    "democratization", "transformative", "groundbreaking", "cutting-edge", "nuanced",
    "multifaceted", "holistic", "paradigm", "paradigms", "unprecedented", "revolutionize",
    "revolutionizes", "revolutionized", "revolutionizing", "harness", "harnesses",
    "harnessed", "harnessing",
]

def _filler_density(text: str) -> float:
    """
    Returns 0.0–1.0 based on filler word density.
    Saturates at 8 per 100 words = 1.0
    """
    words = text.lower().split()
    if not words:
        return 0.0
    count = sum(1 for w in words if w.strip('.,;:!?') in FILLER_WORDS)
    density = (count / len(words)) * 100
    return min(density / 8.0, 1.0)


# ─────────────────────────────────────────────────────────────────────────────
# SIGNAL 6 — Transition phrase overuse
# AI loves "furthermore", "moreover", "in addition", "additionally"
# ─────────────────────────────────────────────────────────────────────────────

TRANSITION_PHRASES = [
    "furthermore", "moreover", "in addition", "additionally", "consequently",
    "therefore", "thus", "hence", "as a result", "subsequently",
    "in conclusion", "to summarize", "in summary", "last but not least",
    "first and foremost", "it is worth noting", "it should be noted",
    "it is important to note", "needless to say",
]

def _transition_overuse(text: str) -> float:
    """
    Returns 0.0–1.0. Saturates at 3 transitions per 100 words.
    """
    word_count = max(len(text.split()), 1)
    lower = text.lower()
    hits = sum(lower.count(t) for t in TRANSITION_PHRASES)
    density = (hits / word_count) * 100
    return min(density / 3.0, 1.0)


# ─────────────────────────────────────────────────────────────────────────────
# AGGREGATE SCORER
# ─────────────────────────────────────────────────────────────────────────────

# Weights for each signal (must sum to 1.0)
WEIGHTS = {
    'ai_phrases':       0.28,
    'burstiness':       0.25,
    'opener_uniformity':0.15,
    'avg_sent_length':  0.12,
    'filler_density':   0.12,
    'transitions':      0.08,
}

_DETECTOR_TOKENIZER = None
_DETECTOR_MODEL = None
_DETECTOR_LOCK = threading.Lock()

_models_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', 'models'))
LOCAL_MODEL_DIR = os.path.join(_models_dir, 'zeal-humanizer-detector-v1.0')

def load_detector():
    """Load the DeBERTa detector model into memory (singleton)."""
    global _DETECTOR_TOKENIZER, _DETECTOR_MODEL
    if not HAS_TORCH_DEPS:
        return None, None

    if _DETECTOR_MODEL is not None:
        return _DETECTOR_TOKENIZER, _DETECTOR_MODEL

    with _DETECTOR_LOCK:
        if _DETECTOR_MODEL is not None:
            return _DETECTOR_TOKENIZER, _DETECTOR_MODEL

        try:
            import torch
        except Exception as e:
            print(f"[Detector] Failed to import torch: {e}")

        # Check local folder first
        if os.path.exists(LOCAL_MODEL_DIR) and os.path.exists(os.path.join(LOCAL_MODEL_DIR, 'model.safetensors')):
            model_path = LOCAL_MODEL_DIR
            print(f"[Detector] Loading desklib model from local folder ({LOCAL_MODEL_DIR})...")
        else:
            model_path = "Zeal000/zeal-humanizer-detector-v1.0"
            print(f"[Detector] Local model not found. Downloading {model_path} from HF Hub...")

        try:
            _DETECTOR_TOKENIZER = AutoTokenizer.from_pretrained(model_path)
            _DETECTOR_MODEL = DesklibAIDetectionModel.from_pretrained(model_path)
            _DETECTOR_MODEL.eval()
            print("[Detector] desklib model loaded successfully.")
        except Exception as e:
            import traceback
            print(f"[Detector] Failed to load DeBERTa detector: {e}")
            traceback.print_exc()

    return _DETECTOR_TOKENIZER, _DETECTOR_MODEL


def _predict_paragraphs_with_deberta(paragraphs: list[str], tokenizer, model) -> tuple[float, list[float]]:
    """Run paragraph-by-paragraph text classification using DeBERTa, returning (overall_prob, paragraph_probs)."""
    import torch
    probabilities = []
    for para in paragraphs:
        encoded = tokenizer(
            para,
            padding='max_length',
            truncation=True,
            max_length=512,
            return_tensors='pt'
        )
        with torch.inference_mode():
            outputs = model(input_ids=encoded['input_ids'], attention_mask=encoded['attention_mask'])
            logits = outputs["logits"]
            logit = logits.item()
            # Calibrate logits: map logit=5.0 to 50%, with a temperature scaling of 1.6
            prob = 1.0 / (1.0 + math.exp(-(logit - 5.0) / 1.6))
            probabilities.append(prob)

    if not probabilities:
        return 0.5, []
    overall = sum(probabilities) / len(probabilities)
    return overall, probabilities


def score_text(text: str, return_chunks: bool = False) -> dict:
    """
    Score a block of text for AI likelihood using fine-tuned DeBERTa model.
    Falls back to the heuristic rule-based scorer if model fails or deps are missing.
    """
    if not text or len(text.split()) < 15:
        return {
            'overall_pct': 0,
            'label': 'Insufficient text',
            'signals': {},
            'word_count': 0,
            'paragraph_count': 0,
            'chunks': [] if return_chunks else None
        }

    # Split text into paragraphs (minimum 8 words to filter out headers/junk in scoring)
    paragraphs = [p.strip() for p in text.split('\n') if len(p.strip().split()) >= 8]
    if not paragraphs:
        paragraphs = [text.strip()]

    tokenizer = None
    model = None
    try:
        tokenizer, model = load_detector()
    except Exception as e:
        print(f"[Detector] Failed to load detector: {e}. Using rule-based fallback.")

    if tokenizer is not None and model is not None:
        try:
            overall_prob, para_probs = _predict_paragraphs_with_deberta(paragraphs, tokenizer, model)
            pct = round(overall_prob * 100)

            # Heuristic signals are computed for UI display
            signals = {
                'ai_phrases':        round(_ai_phrase_density(text) * 100),
                'burstiness':        round(_burstiness_score(text) * 100),
                'opener_uniformity': round(_opener_uniformity(text) * 100),
                'avg_sent_length':   round(_avg_sentence_length_score(text) * 100),
                'filler_density':    round(_filler_density(text) * 100),
                'transitions':       round(_transition_overuse(text) * 100),
            }

            if pct >= 75:
                label = 'Very likely AI-generated'
            elif pct >= 50:
                label = 'Likely AI-generated'
            elif pct >= 30:
                label = 'Possibly AI-assisted'
            elif pct >= 15:
                label = 'Mostly human-sounding'
            else:
                label = 'Reads as human-written'

            res = {
                'overall_pct': pct,
                'label': label,
                'signals': signals,
                'word_count': len(text.split()),
                'paragraph_count': len(paragraphs),
            }
            if return_chunks:
                res['chunks'] = [
                    {'text': p, 'pct': round(prob * 100)}
                    for p, prob in zip(paragraphs, para_probs)
                ]
            return res
        except Exception as e:
            print(f"[Detector] Error running DeBERTa model: {e}. Using rule-based fallback.")

    # Rule-Based Heuristic Scorer Fallback
    # Score each paragraph with rules to get overall and chunk scores
    para_probs = []
    for para in paragraphs:
        para_signals = {
            'ai_phrases':        _ai_phrase_density(para),
            'burstiness':        _burstiness_score(para),
            'opener_uniformity': _opener_uniformity(para),
            'avg_sent_length':   _avg_sentence_length_score(para),
            'filler_density':    _filler_density(para),
            'transitions':       _transition_overuse(para),
        }
        para_weighted = sum(para_signals[k] * WEIGHTS[k] for k in para_signals)
        para_probs.append(para_weighted)

    overall_prob = sum(para_probs) / len(para_probs) if para_probs else 0.5
    pct = round(overall_prob * 100)

    signals = {
        'ai_phrases':        _ai_phrase_density(text),
        'burstiness':        _burstiness_score(text),
        'opener_uniformity': _opener_uniformity(text),
        'avg_sent_length':   _avg_sentence_length_score(text),
        'filler_density':    _filler_density(text),
        'transitions':       _transition_overuse(text),
    }

    if pct >= 75:
        label = 'Very likely AI-generated'
    elif pct >= 50:
        label = 'Likely AI-generated'
    elif pct >= 30:
        label = 'Possibly AI-assisted'
    elif pct >= 15:
        label = 'Mostly human-sounding'
    else:
        label = 'Reads as human-written'

    res = {
        'overall_pct': pct,
        'label': label,
        'signals': {k: round(v * 100) for k, v in signals.items()},
        'word_count': len(text.split()),
        'paragraph_count': len(paragraphs),
    }
    if return_chunks:
        res['chunks'] = [
            {'text': p, 'pct': round(prob * 100)}
            for p, prob in zip(paragraphs, para_probs)
        ]
    return res


# ─────────────────────────────────────────────────────────────────────────────
# DOCX SCORER
# ─────────────────────────────────────────────────────────────────────────────

def score_docx(path: str) -> dict:
    """
    Extract all prose text from a .docx file and score it.
    Returns score dict plus word count and paragraph count.
    """
    from docx import Document

    doc = Document(path)
    prose_blocks = []

    for para in doc.paragraphs:
        t = para.text.strip()
        if t and len(t.split()) >= 5:
            prose_blocks.append(t)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    t = para.text.strip()
                    if t and len(t.split()) >= 5:
                        prose_blocks.append(t)

    full_text = '\n'.join(prose_blocks)
    return score_text(full_text, return_chunks=True)
