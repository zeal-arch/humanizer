# humanizer/pipeline.py
# 17-pass rule-based text humanization pipeline — zero AI API calls

import re
import random
from .phrases import (
    AI_PHRASES,
    INTENSIFIERS,
    DISCOURSE_MARKERS,
    CONTRACTIONS,
    PASSIVE_TO_ACTIVE,
    HEDGE_TRIGGERS,
    PARENTHETICALS,
    SELF_CORRECTIONS,
    PERSONAL_VOICE,
    QUALIFIERS,
    PUNCTUATION_HUMANIZE,
    EM_DASH_REVERSAL,
    PUNCTUATION_HUMANIZE_PROFESSIONAL,
    FRONTING_PATTERNS,
    SYNTACTIC_FRONTING,
    CONJUNCTION_OPENERS,
    ADVERBIAL_FRONTING,
    IDIOMATIC_INJECTIONS,
)
import nltk
from nltk.corpus import wordnet as wn

# ─────────────────────────────────────────────────────────────────────────────
# MODEL REGISTRY — shared singleton so app.py warm-up and pipeline use the
# SAME loaded object (fixes the duplicate-load bug).
# ─────────────────────────────────────────────────────────────────────────────
import os as _os
import threading as _threading

_MODELS_DIR = _os.path.join(_os.path.dirname(__file__), '..', 'models')

QWEN3_DIR   = _os.path.join(_MODELS_DIR, 'Qwen3-0.6b')
T5_DIR      = _os.path.join(_MODELS_DIR, 'T5_Paraphrase_Paws')

_MODEL_TOKENIZER = None
_MODEL_OBJ       = None
_MODEL_TYPE      = None   # 'qwen3' | 't5' | None
_MODEL_LOCK      = _threading.Lock()


def _qwen3_weights_present() -> bool:
    """Check that Qwen3 has its actual weight file, not just config."""
    return (
        _os.path.exists(QWEN3_DIR) and (
            _os.path.exists(_os.path.join(QWEN3_DIR, 'model.safetensors')) or
            _os.path.exists(_os.path.join(QWEN3_DIR, 'pytorch_model.bin'))
        )
    )

def preload_model():
    """
    Load the best available model into the global singleton.
    Called once at startup by app.py — pipeline.py reuses the same object.
    Priority: HF Serverless Inference API (72B) > Qwen3-0.6B local (dev) > HF Hub download (prod) > T5 local (fallback).
    """
    global _MODEL_TOKENIZER, _MODEL_OBJ, _MODEL_TYPE
    if _MODEL_OBJ is not None or _MODEL_TYPE == 'api':
        return  # already loaded

    # ── 1. HF Serverless Inference API (best quality — 72B model, free tier) ──
    # This is the PRIMARY model. The 0.6B local model hallucinates badly,
    # so we always prefer the API when a token is available.
    token = (
        _os.environ.get('HF_API_TOKEN') or
        _os.environ.get('HF_TOKEN') or
        _os.environ.get('humanizeread') or
        _os.environ.get('zeal000') or
        _os.environ.get('HF_READ_TOKEN')
    )
    if token:
        api_model = _os.environ.get('HF_API_MODEL_ID', 'Qwen/Qwen2.5-72B-Instruct')
        print(f"[Model] HF Access Token detected. Initializing Serverless Inference for {api_model}...")
        try:
            from huggingface_hub import InferenceClient
            _MODEL_OBJ = InferenceClient(model=api_model, token=token)
            _MODEL_TYPE = 'api'
            print(f"[Model] HF Serverless Inference client for {api_model} ready.")
            return
        except Exception as e:
            print(f"[Model] Failed to init InferenceClient: {e}. Falling back to local model.")


    # ── 2. Standard Transformers loading ──
    try:
        import torch
        from transformers import AutoTokenizer, AutoModelForCausalLM

        # Optimize dtype: Use bfloat16 or float16 when running on CUDA to minimize VRAM footprint and avoid paging latency,
        # which is extremely critical for 4GB GPUs like the GTX 1650. Fall back to float32 on CPU.
        if torch.cuda.is_available():
            if torch.cuda.is_bf16_supported():
                dtype = torch.bfloat16
            else:
                dtype = torch.float16
        else:
            dtype = torch.float32

        device_kwargs = {
            'device_map': 'auto' if torch.cuda.is_available() else None,
            'dtype': dtype,
        }

        # Enable PyTorch SDPA (Scaled Dot Product Attention) for optimized attention kernels when using GPU
        if torch.cuda.is_available():
            device_kwargs['attn_implementation'] = 'sdpa'

        if _qwen3_weights_present():
            # ── Development: load from local disk ──────────────────────────
            print(f"[Model] Loading Qwen3-0.6B from local disk ({QWEN3_DIR}) ...")
            _MODEL_TOKENIZER = AutoTokenizer.from_pretrained(QWEN3_DIR)
            _MODEL_OBJ = AutoModelForCausalLM.from_pretrained(QWEN3_DIR, **device_kwargs)
            _MODEL_OBJ.eval()
            _MODEL_TYPE = 'qwen3'
            print(f"[Model] Qwen3-0.6B ready on {'GPU' if torch.cuda.is_available() else 'CPU'} (local).")

        else:
            # ── Production: download from Hugging Face Hub ─────────────────
            HF_MODEL_ID = _os.environ.get('HF_MODEL_ID', 'Zeal000/qwen-humanizer-lora')
            print(f"[Model] No local weights found. Downloading {HF_MODEL_ID} from HF Hub ...")
            
            try:
                from peft import PeftConfig, PeftModel
                config = PeftConfig.from_pretrained(HF_MODEL_ID)
                base_model_id = config.base_model_name_or_path
                print(f"[Model] Detected PEFT adapter. Loading base model: {base_model_id}...")
                # Load tokenizer from the BASE model (not the LoRA repo) to avoid
                # a corrupted tokenizer_config.json in the adapter repository.
                _MODEL_TOKENIZER = AutoTokenizer.from_pretrained(base_model_id)
                base_model = AutoModelForCausalLM.from_pretrained(base_model_id, **device_kwargs)
                _MODEL_OBJ = PeftModel.from_pretrained(base_model, HF_MODEL_ID)
                print("[Model] PEFT weights applied successfully.")
            except Exception as peft_e:
                print(f"[Model] Loading as standard model: {peft_e}")
                _MODEL_TOKENIZER = AutoTokenizer.from_pretrained(HF_MODEL_ID)
                _MODEL_OBJ = AutoModelForCausalLM.from_pretrained(HF_MODEL_ID, **device_kwargs)
                
            _MODEL_OBJ.eval()
            _MODEL_TYPE = 'qwen3'
            print(f"[Model] {HF_MODEL_ID} ready on {'GPU' if torch.cuda.is_available() else 'CPU'} (HF Hub).")

    except Exception as e:
        import traceback
        print(f"[Model] Failed to load: {e}")
        traceback.print_exc()


def get_model():
    """Return (tokenizer, model, model_type). Lazy-loads if not already loaded."""
    global _MODEL_TOKENIZER, _MODEL_OBJ, _MODEL_TYPE
    if _MODEL_OBJ is None and _MODEL_TYPE != 'api':
        preload_model()
    return _MODEL_TOKENIZER, _MODEL_OBJ, _MODEL_TYPE


def get_wordnet_pos(treebank_tag):
    if treebank_tag.startswith('J'):
        return wn.ADJ
    elif treebank_tag.startswith('V'):
        return wn.VERB
    elif treebank_tag.startswith('N'):
        return wn.NOUN
    elif treebank_tag.startswith('R'):
        return wn.ADV
    else:
        return None

# Use a truly random seed so every run produces unique output.
# A fixed seed would mean two students running the same AI text
# get identical results — which Turnitin would flag as collusion.
random.seed()


# ─────────────────────────────────────────────────────────────────────────────
# UTILITIES
# ─────────────────────────────────────────────────────────────────────────────

def _replace_dict(text: str, mapping: dict, case_sensitive: bool = False) -> tuple[str, int]:
    """Replace all occurrences of keys with values. Returns (new_text, count)."""
    count = 0
    for src, dst in mapping.items():
        flags = 0 if case_sensitive else re.IGNORECASE
        pattern = re.compile(re.escape(src), flags)

        def _replacer(m, _dst=dst):
            nonlocal count
            count += 1
            original = m.group(0)
            if original and original[0].isupper() and _dst and _dst[0].islower():
                return _dst[0].upper() + _dst[1:]
            return _dst

        text = pattern.sub(_replacer, text)
    return text, count


def _tokenize_sentences(text: str) -> list[str]:
    """Sentence tokenizer using punctuation + capital letter boundaries."""
    parts = re.split(r'(?<=[.!?])\s+(?=[A-Z])', text)
    return [p.strip() for p in parts if p.strip()]


def _split_long_sentence(sentence: str, max_words: int = 24) -> list[str]:
    """Split an overly long sentence at a natural conjunction."""
    words = sentence.split()
    if len(words) <= max_words:
        return [sentence]

    split_words = [
        " but ", " while ",
        " although ", " because ", " since ", " where ", " when ",
        " whereas ", " however ", " so ", " yet ", " unless ",
        " until ", " after ", " before ", " once ", " if ",
    ]

    mid = len(sentence) // 2
    best_pos = -1
    best_dist = len(sentence)
    best_sw = " but "

    for sw in split_words:
        pos = sentence.find(sw, len(sentence) // 4)
        if pos == -1:
            continue
        dist = abs(pos - mid)
        if dist < best_dist:
            best_dist = dist
            best_pos = pos
            best_sw = sw

    if best_pos == -1:
        return [sentence]

    first = sentence[:best_pos].strip().rstrip(',')
    second = sentence[best_pos + len(best_sw):].strip()

    connector_map = {
        " and ": ". Also, ", " but ": ". However, ",
        " while ": ". Meanwhile, ", " although ": ". That said, ",
        " because ": ". The reason is that ", " since ": ". Since ",
        " where ": ". Here ", " when ": ". This happens when ",
        " whereas ": ". In contrast, ", " however ": ". However, ",
        " so ": ". So ", " yet ": ". Yet ",
        " unless ": ". Unless ", " until ": ". Until ",
        " after ": ". After ", " before ": ". Before ",
        " once ": ". Once ", " if ": ". If ",
    }
    prefix = connector_map.get(best_sw, ". ")

    # Lowercase the first letter of second since it will come after prefix words (which are capitalized)
    if second:
        second = second[0].lower() + second[1:]

    # If the prefix is a simple period (starts a new sentence directly), we MUST capitalize second
    if prefix == ". " and second:
        second = second[0].upper() + second[1:]

    if prefix.startswith(". "):
        pfx_text = prefix[2:].strip()
        second = pfx_text + " " + second if pfx_text else second

    return [first + ".", second]


def _merge_short_sentences(sentences: list[str], min_words: int = 7) -> list[str]:
    """Merge consecutive very short sentences with a semicolon."""
    result = []
    i = 0
    while i < len(sentences):
        s = sentences[i]
        if (
            i + 1 < len(sentences)
            and len(s.split()) < min_words
            and len(sentences[i + 1].split()) < min_words
            and random.random() < 0.65
        ):
            next_s = sentences[i + 1].rstrip('.!?')
            merged = s.rstrip('.!?') + "; " + next_s[0].lower() + next_s[1:]
            ending_char = sentences[i + 1][-1]
            if ending_char not in '.!?':
                ending_char = '.'
            merged += ending_char
            result.append(merged)
            i += 2
        else:
            result.append(s)
            i += 1
    return result


def _find_clause_boundary(sentence: str) -> int:
    """Find a good position mid-sentence to inject a parenthetical."""
    # Look for comma positions in the middle half of the sentence
    words = sentence.split()
    if len(words) < 12:
        return -1
    # Find a comma roughly in the middle
    mid = len(sentence) // 2
    for offset in range(0, len(sentence) // 3):
        for sign in [1, -1]:
            pos = mid + sign * offset
            if 0 < pos < len(sentence) and sentence[pos] == ',':
                return pos
    return -1


# ─────────────────────────────────────────────────────────────────────────────
# THE 13 PASSES
# ─────────────────────────────────────────────────────────────────────────────

def pass1_ai_phrases(text: str) -> tuple[str, int]:
    """Replace 500+ known AI phrases with natural equivalents."""
    return _replace_dict(text, AI_PHRASES)


def pass2_intensifiers(text: str) -> tuple[str, int]:
    """Remove or soften intensifiers that AI overuses."""
    return _replace_dict(text, INTENSIFIERS)


def pass3_burstiness(text: str) -> tuple[str, int]:
    """
    Inject sentence length variation — the #1 Turnitin signal.
    Splits long sentences and occasionally merges short consecutive ones.
    Processes paragraphs individually to preserve \n paragraph breaks.
    """
    paragraphs = text.split('\n')
    total_changes = 0
    new_paragraphs = []

    for para in paragraphs:
        stripped = para.strip()
        if not stripped or len(stripped.split()) < 8:
            new_paragraphs.append(para)
            continue

        sentences = _tokenize_sentences(stripped)
        count_before = len(sentences)

        new_sentences = []
        for s in sentences:
            parts = _split_long_sentence(s, max_words=22)
            new_sentences.extend(parts)

        new_sentences = _merge_short_sentences(new_sentences, min_words=9)

        total_changes += abs(len(new_sentences) - count_before)
        new_paragraphs.append(' '.join(new_sentences))

    return '\n'.join(new_paragraphs), total_changes


def pass4_discourse_markers(text: str) -> tuple[str, int]:
    """Prepend human-sounding discourse markers to ~28% of AI-opener paragraphs."""
    ai_openers = (
        "The ", "This ", "These ", "Those ", "In ", "It ", "By ",
        "With ", "As ", "Our ", "We ", "One ", "Another ", "Each ",
        "Such ", "Any ", "Every ", "All ", "Both ", "Many ", "Most ",
    )
    paragraphs = text.split('\n')
    count = 0
    result = []
    used_markers: set[str] = set()

    for para in paragraphs:
        stripped = para.strip()
        if (
            len(stripped.split()) > 12
            and any(stripped.startswith(op) for op in ai_openers)
            and random.random() < 0.28
        ):
            available = [m for m in DISCOURSE_MARKERS if m not in used_markers]
            if not available:
                used_markers.clear()
                available = DISCOURSE_MARKERS[:]

            marker = random.choice(available)
            used_markers.add(marker)

            lowered = stripped[0].lower() + stripped[1:]
            para = para.replace(stripped, marker + lowered, 1)
            count += 1

        result.append(para)

    return '\n'.join(result), count


def pass5_contractions(text: str) -> tuple[str, int]:
    """Inject contractions selectively (skips code-like paragraphs)."""
    if text.count('`') > 2 or text.count('{') > 2 or text.count('//') > 1:
        return text, 0
    return _replace_dict(text, CONTRACTIONS, case_sensitive=True)


def pass6_passive_to_active(text: str) -> tuple[str, int]:
    """Replace weak passive constructions with active voice."""
    return _replace_dict(text, PASSIVE_TO_ACTIVE)


def pass7_opener_diversity(text: str) -> tuple[str, int]:
    """Prevent the same paragraph opener from appearing more than once."""
    replacements = {
        "The system ": ["Our system ", "This system ", "The platform ", "Rayvoy's system "],
        "The project ": ["Our project ", "This project ", "The work ", "The development "],
        "The platform ": ["Our platform ", "This platform ", "The system ", "The tool "],
        "The proposed ": ["Our proposed ", "The suggested ", "This proposed "],
        "The existing ": ["Our existing ", "The current ", "The old ", "The prior "],
        "The new ": ["Our new ", "The updated ", "This new ", "The revised "],
        "The team ": ["Our team ", "The staff ", "The group ", "Our counsellors "],
        "The user ": ["Each user ", "A user ", "The counsellor ", "Any user "],
        "The data ": ["All data ", "This data ", "Student data ", "The records "],
        "The process ": ["This process ", "Our process ", "The workflow ", "The procedure "],
        "The system is ": ["Our system is ", "The platform is ", "This system is "],
        "The approach ": ["Our approach ", "This approach ", "The method "],
        "The solution ": ["Our solution ", "This solution ", "The tool "],
        "The feature ": ["This feature ", "Our feature ", "Each feature "],
        "The module ": ["This module ", "Our module ", "Each module "],
        "The database ": ["Our database ", "This database ", "The backend "],
        "The interface ": ["Our interface ", "This interface ", "The UI "],
        "The implementation ": ["Our implementation ", "This implementation "],
    }

    paragraphs = text.split('\n')
    opener_counts: dict[str, int] = {}
    count = 0
    result = []

    for para in paragraphs:
        stripped = para.strip()
        for opener, alternatives in replacements.items():
            if stripped.startswith(opener):
                opener_counts[opener] = opener_counts.get(opener, 0) + 1
                if opener_counts[opener] >= 2:
                    alt = random.choice(alternatives)
                    para = para.replace(opener, alt, 1)
                    count += 1
                break
        result.append(para)

    return '\n'.join(result), count


def pass8_hedging(text: str) -> tuple[str, int]:
    """
    Replace overconfident AI assertions with hedged human-like language.
    Humans naturally hedge — AI is always certain.
    """
    return _replace_dict(text, HEDGE_TRIGGERS, case_sensitive=False)


def pass9_parentheticals(text: str) -> tuple[str, int]:
    """
    Inject parenthetical asides into ~15% of long sentences.
    Humans add tangential remarks mid-sentence — AI never does.
    """
    sentences = _tokenize_sentences(text)
    count = 0
    new_sentences = []
    used: set[str] = set()

    word_count = sum(len(s.split()) for s in sentences)
    multiplier = 3.0 if word_count < 100 else 2.0 if word_count < 200 else 1.0
    for s in sentences:
        words = s.split()
        if len(words) >= 18 and random.random() < (0.08 * multiplier):
            # Find a comma to inject after, or inject before the last clause
            boundary = _find_clause_boundary(s)
            available = [p for p in PARENTHETICALS if p not in used]
            if not available:
                used.clear()
                available = PARENTHETICALS[:]

            aside = random.choice(available)
            used.add(aside)

            if boundary > 0:
                new_s = s[:boundary + 1] + aside + s[boundary + 1:]
            else:
                # Inject before the last 4 words
                word_list = s.split()
                split_at = max(len(word_list) - 4, len(word_list) // 2)
                new_s = ' '.join(word_list[:split_at]) + aside + ' ' + ' '.join(word_list[split_at:])

            new_sentences.append(new_s)
            count += 1
        else:
            new_sentences.append(s)

    return ' '.join(new_sentences), count


def pass10_self_corrections(text: str) -> tuple[str, int]:
    """
    Insert human-style self-correction patterns at ~10% of sentence boundaries.
    'Or rather,', 'More precisely,', 'To be specific,' — AI never self-corrects.
    """
    count = 0
    word_count = len(text.split())
    multiplier = 3.0 if word_count < 100 else 2.0 if word_count < 200 else 1.0
    for src, dst in SELF_CORRECTIONS.items():
        # Only apply with a probability per occurrence
        occurrences = [m.start() for m in re.finditer(re.escape(src), text)]
        for pos in sorted(occurrences, reverse=True):
            if random.random() < (0.15 * multiplier):
                text = text[:pos] + dst + text[pos + len(src):]
                count += 1
    return text, count


def pass11_personal_voice(text: str) -> tuple[str, int]:
    """
    Replace impersonal 'one can' / 'it can be' with personal voice.
    Human writers use 'we found', 'we noticed' — AI stays impersonal.
    """
    return _replace_dict(text, PERSONAL_VOICE)


def pass12_qualifiers(text: str) -> tuple[str, int]:
    """
    Replace absolute AI statements with qualified human ones.
    AI says 'always' and 'completely' — humans say 'usually' and 'largely'.
    """
    return _replace_dict(text, QUALIFIERS)


def pass13_punctuation(text: str, register: str = 'casual') -> tuple[str, int]:
    """
    Humanize punctuation patterns: replace bland conjunctions with
    semicolons and varied connectors (explicitly avoiding em-dashes).
    """
    if text.count('`') > 2 or text.count('{') > 2:
        return text, 0
    if register == 'professional':
        return _replace_dict(text, PUNCTUATION_HUMANIZE_PROFESSIONAL, case_sensitive=True)
    return _replace_dict(text, PUNCTUATION_HUMANIZE, case_sensitive=True)


def pass14_syntactic_fronting(text: str) -> tuple[str, int]:
    """
    Break up repetitive Subject-Verb-Object structures.
    Uses regex for complex clefting and dict replacement for simple fronting.
    """
    count = 0
    
    # 1. Regex-based fronting
    for pattern, replacement in FRONTING_PATTERNS:
        matches = len(re.findall(pattern, text))
        if matches > 0:
            text = re.sub(pattern, replacement, text)
            count += matches
            
    # 2. Simple dictionary fronting
    text, n1 = _replace_dict(text, SYNTACTIC_FRONTING)
    text, n2 = _replace_dict(text, ADVERBIAL_FRONTING)
    
    return text, count + n1 + n2


def pass15_perplexity_guided_perturbation(text: str) -> tuple[str, int]:
    """
    Perplexity-guided word perturbation — the core anti-GPTZero pass.

    Uses GPT-2 to score each sentence's perplexity (predictability).
    Sentences with LOW perplexity (< 35) are "too AI-like" — GPTZero
    will flag them. We aggressively perturb those sentences with WordNet
    synonym rotation, filtered through student vocabulary.

    Sentences with HIGH perplexity (> 35) are already "human-enough"
    and are left untouched to preserve quality.

    This directly attacks what GPTZero measures: the token probability
    distribution. By replacing predictable tokens with less predictable
    alternatives, we raise per-sentence perplexity above the detection
    threshold.
    """
    from .perplexity import sentence_perplexity, PPL_AI_THRESHOLD, find_predictable_tokens, PPL_MAX_ITERATIONS
    from .student_vocab import STUDENT_VOCABULARY, filter_student_synonyms
    from nltk.corpus import words as nltk_words

    try:
        common_words = set(w.lower() for w in nltk_words.words())
    except Exception:
        common_words = set()

    # Banned synonyms (WordNet garbage and awkward/informal synonyms)
    BANNED = {
        'nidus', 'clobber', 'solemnisation', 'solemnization', 'clip', 'cartridge',
        'knockout', 'telling', 'tellings', 'animation', 'animations', 'fasting',
        'eternal', 'perpetual', 'tempt', 'mold', 'mould', 'forge', 'fashion',
        'plication', 'plica', 'crease', 'crinkle', 'rumple', 'cockle',
        'piss', 'urinate', 'piddle', 'micturate', 'wee', 'defecate',
        'crap', 'stool', 'slew', 'slay', 'smite', 'smote',
        'beget', 'engender', 'sire', 'spawn', 'hump', 'bonk', 'copulate',
        'fornicate', 'intercourse', 'coitus', 'congress',
        'snuff', 'croak', 'decease', 'perish', 'conk',
        'stratum', 'substrate', 'substratum', 'substructure',
        'bum', 'hobo', 'tramp', 'vagrant', 'mooch', 'panhandle',
        'incur', 'sustain', 'brook', 'stomach', 'abide', 'endure',
        'ilk', 'kidney', 'kinsfolk', 'kinfolk', 'kinship',
        'phallus', 'member', 'penis', 'privy', 'commode', 'throne',
        'dolt', 'dullard', 'dunce', 'bonehead', 'numskull', 'blockhead',
        'boob', 'booby', 'nincompoop', 'ninny', 'simpleton',
        'stub', 'butt', 'rump', 'rear', 'derriere', 'fundament',
        'bugger', 'blighter', 'chap', 'cuss', 'fellow', 'geezer',
        'elds', 'agnise', 'agnize', 'cosmos', 'macrocosm',

        # Awkward/informal synonyms that AI or spin-bots introduce
        'peeps', 'peep', 'confab', 'confabbing', 'confabbed', 'say', 'dismay',
        'gos', 'decree', 'enquiry', 'enquiries', 'gathering', 'gatherings',
        'growing', 'creation', 'creations', 'living', 'citizenry', 'multitude',
        'arouse', 'companion', 'labor', 'wont', 'drill', 'space', 'peculiarly',
        'pass', 's', 'patch', 'watershed', 'conflict', 'equilibrium', 'chats',
        'chat', 'alarm', 'alarms', 'alarmed', 'instruction', 'instructions',
        'occurred', 'aroused', 'head', 'hand', 'eye', 'heart', 'face', 'body',
        'office', 'turn', 'turns', 'step', 'steps', 'part', 'parts', 'course', 'courses'
    }

    SKIP_WORDS = {
        'the', 'a', 'an', 'is', 'are', 'was', 'were', 'be', 'been', 'being',
        'have', 'has', 'had', 'do', 'does', 'did', 'will', 'would', 'could',
        'should', 'may', 'might', 'shall', 'can', 'must', 'need',
        'i', 'you', 'he', 'she', 'it', 'we', 'they', 'me', 'him', 'her',
        'us', 'them', 'my', 'your', 'his', 'its', 'our', 'their',
        'this', 'that', 'these', 'those', 'what', 'which', 'who', 'whom',
        'not', 'no', 'yes', 'so', 'if', 'or', 'and', 'but', 'for', 'nor',
        'on', 'at', 'to', 'by', 'in', 'of', 'up', 'out', 'off', 'over',
        'into', 'with', 'from', 'about', 'than', 'then', 'now', 'just',
        'also', 'very', 'too', 'much', 'more', 'most', 'own', 'same',
        'other', 'each', 'every', 'both', 'few', 'many', 'some', 'any',
        'such', 'only', 'even', 'still', 'already', 'always', 'never',
        'get', 'got', 'go', 'went', 'gone', 'come', 'came', 'make', 'made',
        'take', 'took', 'give', 'gave', 'say', 'said', 'tell', 'told',
        'know', 'knew', 'think', 'thought', 'see', 'saw', 'want', 'like',
        'use', 'used', 'find', 'found', 'put', 'set', 'run', 'let',
        'really', 'pretty', 'quite', 'kinda', 'stuff', 'things', 'thing',
        'look', 'looks', 'feel', 'feels', 'seem', 'seems', 'keep', 'keeps',
        'turn', 'turns', 'show', 'shows', 'call', 'calls', 'help', 'helps',
        'work', 'works', 'play', 'plays', 'live', 'lives', 'move', 'moves',
        'long', 'back', 'down', 'well', 'good', 'hard', 'fast', 'kind',
        'able', 'real', 'sure', 'open', 'full', 'late', 'easy', 'high',
        'best', 'last', 'next', 'near', 'free', 'left', 'right',
        'don', 'doesn', 'didn', 'won', 'wouldn', 'couldn', 'shouldn',
        't', 's', 're', 've', 'll', 'd', 'm',
        'track', 'course', 'lead', 'step', 'steps', 'life', 'lives', 'spot',
        'role', 'roles', 'routine', 'routines', 'system', 'systems', 'process',
        'processes', 'method', 'methods', 'development', 'developments', 'progress',
        'consistency', 'motivation', 'productivity', 'behavior', 'behaviors', 'pattern',
        'patterns', 'habit', 'habits', 'action', 'actions', 'outcome', 'outcomes',
        'result', 'results', 'change', 'changes', 'experience', 'experiences',
        'memory', 'memories', 'city', 'cities', 'neighborhood', 'neighborhoods',
        'structure', 'structures', 'space', 'spaces', 'street', 'streets',
        'landmark', 'landmarks', 'preservation', 'progress', 'decision', 'decisions',
        'practice', 'practices', 'discipline', 'disciplines', 'growth', 'grow'
    }

    def _double_consonant(word: str) -> str:
        if len(word) >= 3 and word[-1] in 'bdglmnprt' and word[-2] in 'aeiou' and word[-3] not in 'aeiou':
            return word + word[-1]
        return word

    def _get_synonym(word: str, tag: str) -> str | None:
        """Try to find a student-level synonym for a word, preserving grammar."""
        trailing = ''
        core = word
        while core and core[-1] in '.,;:!?\'")-':
            trailing = core[-1] + trailing
            core = core[:-1]

        clean = re.sub(r'[^a-zA-Z]', '', core)
        wn_pos = get_wordnet_pos(tag)

        if (wn_pos is None or wn_pos == wn.NOUN or len(clean) < 4 or clean.lower() in SKIP_WORDS
                or tag in ('NNP', 'NNPS')):
            return None

        synsets = wn.synsets(clean.lower(), pos=wn_pos)
        if not synsets:
            return None

        # Check first 2 synsets for more candidates
        candidates = set()
        for syn in synsets[:2]:
            for lemma in syn.lemmas():
                name = lemma.name().replace('_', ' ')
                if (' ' not in name
                    and name.lower() != clean.lower()
                    and abs(len(name) - len(clean)) <= 4
                    and name.lower() not in BANNED
                    and (not common_words or name.lower() in common_words)):
                    candidates.add(name)

        if not candidates:
            return None

        # Prefer student vocabulary synonyms
        student_filtered = filter_student_synonyms(candidates, clean)
        if not student_filtered:
            return None

        synonym = random.choice(student_filtered)

        # Inflection rules
        # 1. Plural nouns (NNS)
        if tag == 'NNS' or (tag == 'NN' and clean.endswith('s') and not clean.endswith('ss')):
            if not synonym.endswith('s'):
                if synonym.endswith('y') and not any(synonym.endswith(x) for x in ['ay', 'ey', 'oy', 'uy']):
                    synonym = synonym[:-1] + 'ies'
                elif any(synonym.endswith(x) for x in ['ch', 'sh', 'x', 's', 'z']):
                    synonym += 'es'
                else:
                    synonym += 's'
        # 2. Gerunds / present participles (VBG)
        elif tag == 'VBG':
            if not synonym.endswith('ing'):
                synonym = _double_consonant(synonym)
                if synonym.endswith('e') and not any(synonym.endswith(x) for x in ['ee', 'oe', 'ye']):
                    synonym = synonym[:-1] + 'ing'
                elif synonym.endswith('ie'):
                    synonym = synonym[:-2] + 'ying'
                else:
                    synonym += 'ing'
        # 3. Past tense / past participles (VBD, VBN)
        elif tag in ('VBD', 'VBN'):
            IRREGULAR_PAST = {
                'understand': 'understood', 'build': 'built', 'spend': 'spent',
                'keep': 'kept', 'meet': 'met', 'leave': 'left', 'lose': 'lost',
                'sell': 'sold', 'tell': 'told', 'make': 'made', 'hold': 'held',
                'bring': 'brought', 'think': 'thought', 'feel': 'felt',
                'choose': 'chosen', 'see': 'seen', 'know': 'known', 'write': 'written',
                'take': 'taken', 'give': 'given', 'do': 'done', 'run': 'run',
                'grow': 'grown', 'find': 'found', 'get': 'got', 'go': 'went',
                'say': 'said', 'hear': 'heard', 'rise': 'risen', 'speak': 'spoken',
                'fall': 'fallen', 'become': 'became', 'break': 'broken', 'bend': 'bent',
                'catch': 'caught', 'draw': 'drawn', 'drink': 'drunk', 'drive': 'driven',
                'eat': 'eaten', 'forget': 'forgotten', 'forgive': 'forgiven', 'hide': 'hidden',
                'ride': 'ridden', 'ring': 'rung', 'run': 'run', 'shake': 'shaken',
                'sing': 'sung', 'sink': 'sunk', 'steal': 'stolen', 'strike': 'struck',
                'swear': 'sworn', 'swim': 'swum', 'tear': 'torn', 'throw': 'thrown',
                'wear': 'worn', 'win': 'won'
            }
            if synonym in IRREGULAR_PAST:
                synonym = IRREGULAR_PAST[synonym]
            elif not synonym.endswith('ed'):
                synonym = _double_consonant(synonym)
                if synonym.endswith('e'):
                    synonym += 'd'
                elif synonym.endswith('y') and not any(synonym.endswith(x) for x in ['ay', 'ey', 'oy', 'uy']):
                    synonym = synonym[:-1] + 'ied'
                else:
                    synonym += 'ed'
        # 4. 3rd person singular present verbs (VBZ)
        elif tag == 'VBZ':
            if not synonym.endswith('s'):
                if synonym.endswith('y') and not any(synonym.endswith(x) for x in ['ay', 'ey', 'oy', 'uy']):
                    synonym = synonym[:-1] + 'ies'
                elif any(synonym.endswith(x) for x in ['ch', 'sh', 'x', 's', 'z']):
                    synonym += 'es'
                else:
                    synonym += 's'
        # 5. Adverbs ending in ly
        elif tag == 'RB' and clean.endswith('ly'):
            NON_LY_ADVERBS = {
                'always', 'never', 'often', 'sometimes', 'soon', 'fast', 'here', 'there',
                'now', 'then', 'where', 'why', 'how', 'far', 'hard', 'late', 'well',
                'already', 'still', 'also', 'seldom', 'rather', 'quite', 'too', 'very',
                'even', 'almost', 'maybe', 'perhaps', 'instead', 'otherwise', 'therefore',
                'however', 'together', 'apart', 'alone', 'yet', 'more', 'most', 'less', 'least',
                'much', 'little'
            }
            if synonym.lower() not in NON_LY_ADVERBS and not synonym.endswith('ly'):
                if synonym.endswith('y'):
                    synonym = synonym[:-1] + 'ily'
                else:
                    synonym += 'ly'

        # Preserve capitalization
        if core and core[0].isupper():
            synonym = synonym[0].upper() + synonym[1:]

        return synonym + trailing

    # Process paragraph-by-paragraph to PRESERVE paragraph breaks
    paragraphs = text.split('\n')
    total_changes = 0
    new_paragraphs = []

    for para in paragraphs:
        stripped = para.strip()
        if not stripped or len(stripped.split()) < 8:
            new_paragraphs.append(para)
            continue

        # Split paragraph into sentences
        sentences = _tokenize_sentences(stripped)
        new_sentences = []

        for sent in sentences:
            words = sent.split()
            if len(words) < 5:
                new_sentences.append(sent)
                continue

            # Score this sentence's perplexity
            ppl = sentence_perplexity(sent)

            if ppl >= PPL_AI_THRESHOLD:
                # Already "human enough" — don't touch it
                new_sentences.append(sent)
                continue

            # This sentence is "too AI" — try up to 3 rounds of targeted perturbation
            current_sent = sent
            current_ppl = ppl
            round_changes = 0

            for attempt in range(PPL_MAX_ITERATIONS):
                if current_ppl >= PPL_AI_THRESHOLD:
                    break

                # Find the most predictable tokens in the current sentence
                predictable = find_predictable_tokens(current_sent, top_n=5)
                if not predictable:
                    break

                predictable_words = set(tok.strip().lower() for tok, prob, pos in predictable)

                current_words = current_sent.split()
                try:
                    tagged = nltk.pos_tag(current_words)
                except Exception:
                    break

                new_words = []
                attempt_changes = 0
                for word, tag in tagged:
                    clean_w = re.sub(r'[^a-zA-Z]', '', word).lower()
                    
                    # Target only if the word matches or contains a predictable token
                    is_predictable = clean_w in predictable_words or any(tok in clean_w for tok in predictable_words if len(tok) >= 3)
                    
                    # Give a small random chance (15%) to perturb other content words for sentence-level variety
                    if is_predictable or (random.random() < 0.15 and len(clean_w) >= 4 and clean_w not in SKIP_WORDS):
                        replacement = _get_synonym(word, tag)
                        if replacement:
                            new_words.append(replacement)
                            attempt_changes += 1
                            round_changes += 1
                            total_changes += 1
                        else:
                            new_words.append(word)
                    else:
                        new_words.append(word)

                if attempt_changes == 0:
                    break  # No more words to swap

                current_sent = ' '.join(new_words)
                current_ppl = sentence_perplexity(current_sent)

            new_sentences.append(current_sent)

            if round_changes > 0:
                print(f"  [pass15] PPL {ppl:.0f} -> {current_ppl:.0f} "
                      f"({round_changes} swaps) | {sent[:50]}...")

        new_paragraphs.append(' '.join(new_sentences))

    return '\n'.join(new_paragraphs), total_changes


def pass16_imperfect_discourse(text: str) -> tuple[str, int]:
    """
    Inject conjunction openers (And, But, Yet) to start some paragraphs.
    AI is trained to avoid these, making them a strong human signal.
    """
    paragraphs = text.split('\n')
    count = 0
    result = []
    
    word_count = sum(len(p.split()) for p in paragraphs)
    multiplier = 3.0 if word_count < 100 else 2.0 if word_count < 200 else 1.0
    
    for para in paragraphs:
        stripped = para.strip()
        # Avoid stacking if it already starts with a conjunction or a pass4 marker
        existing_markers = tuple(list(CONJUNCTION_OPENERS) + list(DISCOURSE_MARKERS))
        if (
            len(stripped.split()) > 15
            and not any(stripped.startswith(op.strip()) for op in existing_markers)
            and random.random() < (0.08 * multiplier)
        ):
            opener = random.choice(CONJUNCTION_OPENERS)
            
            # Lowercase the original first word if it's not 'I' or an acronym
            first_word_end = stripped.find(' ')
            if first_word_end > 0:
                first_word = stripped[:first_word_end]
                if first_word != "I" and not first_word.isupper():
                    lowered = first_word.lower() + stripped[first_word_end:]
                    para = para.replace(stripped, opener + lowered, 1)
                else:
                    para = para.replace(stripped, opener + stripped, 1)
            count += 1
            
        result.append(para)
        
    return '\n'.join(result), count


def _rewrite_with_inference_api(para: str, client, register: str = 'casual') -> str:
    """
    Use Hugging Face Serverless Inference API (InferenceClient) with a 72B model.
    This is the PRIMARY rewrite engine — produces human-quality output.
    """
    if register == 'professional':
        prompt = (
            f"Rewrite this paragraph to sound like a real person wrote it for a professional document. "
            f"Keep the exact same facts and meaning. Use clear, direct language. "
            f"Do NOT use AI words like 'delve', 'leverage', 'robust', 'crucial', 'furthermore'. "
            f"Do NOT use em dashes. Mix sentence lengths naturally. "
            f"Output ONLY the rewritten paragraph.\n\n"
            f"{para}"
        )
        system_content = "You rewrite text to sound naturally human. Output only the rewritten text, nothing else."
    else:
        prompt = (
            f"Rewrite this paragraph so it sounds like a real person casually wrote it. "
            f"Keep ALL the same events, details, and meaning. Don't add new information. "
            f"Don't remove any key details. "
            f"Use contractions (I'm, didn't, it's, wasn't). Use simple everyday words. "
            f"Vary sentence length, some short, some longer. "
            f"Do NOT use fancy vocabulary or AI-sounding words. "
            f"Output ONLY the rewritten paragraph, nothing else.\n\n"
            f"{para}"
        )
        system_content = "You rewrite text to sound naturally human and casual. Output only the rewritten text, nothing else. Never explain what you did."

    messages = [
        {"role": "system", "content": system_content},
        {"role": "user", "content": prompt}
    ]

    # Cap output length to ~1.5x input to prevent rambling
    max_tokens = min(400, int(len(para.split()) * 2.0) + 40)

    try:
        response = client.chat_completion(
            messages=messages,
            max_tokens=max_tokens,
            temperature=0.85,
            top_p=0.90,
        )
    except Exception as e:
        print(f"[API] Inference API error: {e}")
        return para  # return original on API failure

    generated = response.choices[0].message.content.strip()

    # Clean up model artifacts
    generated = re.sub(r'<think>.*?</think>', '', generated, flags=re.DOTALL).strip()
    generated = re.sub(r'^(Rewritten|Output|Here is|Here\'s|Result|Sure)[:\s]+', '', generated, flags=re.IGNORECASE).strip()
    # Strip wrapping quotes if the model quoted the whole output
    if len(generated) > 2 and generated[0] == '"' and generated[-1] == '"':
        generated = generated[1:-1].strip()
    generated = generated.replace('\u2018', "'").replace('\u2019', "'").replace('\u201c', '"').replace('\u201d', '"')
    generated = re.sub(r'\s+', ' ', generated)
    generated = re.sub(r'\s+([.,!?;:])', r'\1', generated)

    words_out = len(generated.split())
    words_in  = len(para.split())
    if words_out < 3 or words_out > words_in * 2.5:
        return para  # safety fallback — reject if output is way too long/short

    return generated


def _has_hallucination(original: str, generated: str) -> bool:
    """
    Check if the generated text has hallucinated content not in the original.
    A small model (0.6B) often invents completely new sentences, characters,
    or scenarios. We detect this by checking if the generated text introduces
    question marks, quotation patterns, or dramatically different content.
    """
    # If output has way more sentences than input, likely hallucinating
    orig_sentences = len(re.split(r'[.!?]+', original))
    gen_sentences = len(re.split(r'[.!?]+', generated))
    if gen_sentences > orig_sentences * 2.5:
        return True

    # If original has no questions but output has multiple, it's hallucinating
    orig_questions = original.count('?')
    gen_questions = generated.count('?')
    if orig_questions == 0 and gen_questions >= 2:
        return True

    # If the output introduces "Original sentence:" or meta-commentary
    meta_patterns = [
        r'original\s+(sentence|text|paragraph)',
        r'rewritten\s+(sentence|text|paragraph)',
        r'reread\s+carefully',
        r'let\'s\s+clarify',
        r'how\s+did\s+he\s+know',
        r'isn\'t\s+he\s+waiting',
    ]
    for pattern in meta_patterns:
        if re.search(pattern, generated, re.IGNORECASE):
            return True

    return False


def _rewrite_with_qwen3(para: str, tokenizer, model, register: str = 'casual') -> str:
    """
    Use Qwen3-0.6B as a FALLBACK rewriter when the API is unavailable.
    WARNING: This model is tiny (620M params) and prone to hallucination.
    We use very conservative settings and strict output validation.
    """
    import torch

    # Keep the prompt extremely simple for the tiny model
    if register == 'professional':
        prompt = (
            f"Rewrite this paragraph in a professional tone. Keep the same meaning. "
            f"Output only the rewritten text.\n\n{para}"
        )
        system_content = "Rewrite text professionally. Output only the rewritten text."
    else:
        prompt = (
            f"Rewrite this paragraph casually. Keep the same meaning and all details. "
            f"Use contractions. Output only the rewritten text.\n\n{para}"
        )
        system_content = "Rewrite text casually. Output only the rewritten text."

    messages = [
        {"role": "system", "content": system_content},
        {"role": "user", "content": prompt}
    ]

    # Apply chat template — enable_thinking=False disables Qwen3's <think> block
    try:
        encoded = tokenizer.apply_chat_template(
            messages,
            add_generation_prompt=True,
            return_tensors="pt",
            enable_thinking=False,
        )
        input_ids = encoded["input_ids"] if hasattr(encoded, "keys") else encoded
    except TypeError:
        try:
            encoded = tokenizer.apply_chat_template(
                messages,
                add_generation_prompt=True,
                return_tensors="pt",
            )
            input_ids = encoded["input_ids"] if hasattr(encoded, "keys") else encoded
        except Exception:
            encoded = tokenizer.encode(prompt, return_tensors="pt")
            input_ids = encoded["input_ids"] if hasattr(encoded, "keys") else encoded
    except Exception:
        encoded = tokenizer.encode(prompt, return_tensors="pt")
        input_ids = encoded["input_ids"] if hasattr(encoded, "keys") else encoded

    if isinstance(input_ids, list):
        input_ids = torch.tensor([input_ids])

    input_ids = input_ids.to(model.device)
    input_len = input_ids.shape[-1]
    attention_mask = torch.ones_like(input_ids).to(model.device)

    # Conservative token cap — don't let the tiny model ramble
    max_tokens = min(250, int(len(para.split()) * 1.8) + 30)

    with _MODEL_LOCK:
        with torch.inference_mode():
            outputs = model.generate(
                input_ids,
                attention_mask=attention_mask,
                max_new_tokens=max_tokens,
                do_sample=True,
                temperature=0.70,   # Lower temp = less hallucination
                top_p=0.85,
                top_k=30,
                repetition_penalty=1.2,
                use_cache=True,
                pad_token_id=tokenizer.eos_token_id,
            )

    generated = tokenizer.decode(outputs[0][input_len:], skip_special_tokens=True).strip()

    generated = re.sub(r'<think>.*?</think>', '', generated, flags=re.DOTALL).strip()
    # Strip any "Rewritten:" / "Output:" label the model might prepend
    generated = re.sub(r'^(Rewritten|Output|Here is|Here\'s|Result)[:\s]+', '', generated, flags=re.IGNORECASE).strip()

    # Standardize quotation marks and curly apostrophes to straight ones for clean encoding and maximum human likeness
    generated = generated.replace('’', "'").replace('‘', "'").replace('“', '"').replace('”', '"')
    # Clean up double spaces or spaces before punctuation
    generated = re.sub(r'\s+', ' ', generated)
    generated = re.sub(r'\s+([.,!?;:])', r'\1', generated)

    words_out = len(generated.split())
    words_in  = len(para.split())

    # Strict safety: tiny model often rambles or hallucinates
    if words_out < 3 or words_out > words_in * 2.5:
        print("[Qwen3] Rejected: length mismatch (in={}, out={})".format(words_in, words_out))
        return para

    if _has_hallucination(para, generated):
        print("[Qwen3] Rejected: hallucination detected")
        return para

    return generated


def _rewrite_with_t5(para: str, tokenizer, model) -> str:
    """T5 fallback using nucleus sampling for high perplexity."""
    import torch
    text_input = "paraphrase: " + para + " </s>"
    input_ids = tokenizer.encode(text_input, return_tensors="pt", max_length=256, truncation=True)
    input_ids = input_ids.to(model.device)

    # Acquire model lock to guarantee thread-safe autoregressive decoding (prevents text bleed)
    with _MODEL_LOCK:
        # Inference mode is faster and uses less VRAM/memory than no_grad
        with torch.inference_mode():
            outputs = model.generate(
                input_ids,
                max_length=256,
                do_sample=True,
                top_k=50,
                top_p=0.95,
                temperature=1.2,
                num_return_sequences=1,
                no_repeat_ngram_size=2,
                use_cache=True,
            )

    out_text = tokenizer.decode(outputs[0], skip_special_tokens=True)
    out_text = re.sub(r'\s+([.,!?;:])', r'\1', out_text)
    return out_text


def pass19_structural_smoothing(text: str, progress_callback=None, register: str = 'casual') -> tuple[str, int]:
    """
    Neural rewrite pass — uses the best available model (Serverless API or local).
    HF Serverless API (preferred): state-of-the-art 72B model.
    Qwen3-0.6B / 3B (local/Hub): instruction-tuned causal LM.
    T5_Paraphrase_Paws (fallback): nucleus sampling paraphrase.
    """
    tokenizer, model, model_type = get_model()
    if model is None:
        return text, 0
    if model_type != 'api' and tokenizer is None:
        return text, 0

    paragraphs = text.split('\n')
    smoothed_paragraphs = []
    changes = 0
    total = len(paragraphs)

    for i, para in enumerate(paragraphs):
        stripped = para.strip()

        # Optimize: skip the expensive neural model for short sentences/paragraphs under 12 words.
        # Rule-based passes will still humanize them instantly.
        if not stripped or len(stripped.split()) < 12:
            smoothed_paragraphs.append(para)
            continue

        if progress_callback:
            model_label = "HF API 72B" if model_type == 'api' else ("Qwen3" if model_type == 'qwen3' else "T5")
            progress_callback(i, total, f"[{model_label}] Rewriting paragraph {i+1} of {total}...")

        try:
            if model_type == 'api':
                out_text = _rewrite_with_inference_api(stripped, model, register=register)
            elif model_type == 'qwen3':
                out_text = _rewrite_with_qwen3(stripped, tokenizer, model, register=register)
            else:
                out_text = _rewrite_with_t5(stripped, tokenizer, model)

            if out_text.lower() != stripped.lower():
                changes += 1
            smoothed_paragraphs.append(out_text)
        except Exception as e:
            print(f"[pass19] Error on paragraph {i+1}: {e}")
            smoothed_paragraphs.append(para)  # keep original on error

    return '\n'.join(smoothed_paragraphs), changes


def pass17_ghost_characters(text: str) -> tuple[str, int]:
    """
    Remove invisible zero-width spaces often left by copying from ChatGPT.
    """
    ghost_chars = ['\u200b', '\u200c', '\u200d', '\ufeff']
    count = 0
    for gc in ghost_chars:
        if gc in text:
            count += text.count(gc)
            text = text.replace(gc, '')
    return text, count


def pass20_em_dash_reversal(text: str) -> tuple[str, int]:
    """
    Strip AI-introduced em-dashes (—) from the neural rewrite output.

    Em-dash overuse is one of the most reliable AI-writing signals:
    - Wikipedia:Signs of AI writing #4.4: "Overuse of em dashes"
    - SKILL.md #14: AI uses em dashes ~10x more than humans
    - Ghostwriter: listed as a top mechanical tell

    This pass runs AFTER the neural rewrite (pass19) because that's where
    most em-dashes get introduced. It replaces 'X — Y' patterns with
    simpler commas, which is what human writers naturally use.
    """
    return _replace_dict(text, EM_DASH_REVERSAL, case_sensitive=True)

def pass18_perplexity_tension(text: str, register: str = 'casual') -> tuple[str, int]:
    """
    Target AIW-2 model by shattering next-word predictability (perplexity).
    Randomly injects organic idioms and relies on NLTK for the rest.
    """
    if register == 'professional':
        return text, 0
    text, changes_idioms = _replace_dict(text, {k: random.choice(v) for k, v in IDIOMATIC_INJECTIONS.items()})
    
    # Rare synonym injection is now handled universally by NLTK in Pass 15
    return text, changes_idioms


# ─────────────────────────────────────────────────────────────────────────────
# MAIN PIPELINE ORCHESTRATOR
# ─────────────────────────────────────────────────────────────────────────────

def humanize_text(text: str, progress_callback=None, register: str = 'casual') -> dict:
    """
    Run the humanization pipeline.

    Strategy to defeat GPTZero Model 4.6b and neural AI classifiers:
    1. Neural rewrite runs FIRST on clean text (72B API or local Qwen3/T5).
       This produces fluent rewrites but the statistical token distribution
       is still detectable by neural classifiers like DeBERTa and GPTZero.
    2. Structural perturbation passes run AFTER the neural rewrite to break
       the token-level probability patterns that classifiers detect:
       - Burstiness: split/merge sentences to vary rhythm
       - Punctuation: em-dashes/semicolons change token boundaries
       - Syntactic fronting: move clauses to break SVO patterns
       - Passive-to-active: structural transformation
       - Imperfect discourse: conjunction openers (human-only signal)
    3. Light cleanup passes polish the output:
       - AI phrase removal, contractions, hedging, intensifier softening
    4. Heavy fingerprint passes (parentheticals, self-corrections) stay
       DISABLED to avoid the humanizer-tool detection heuristic.
    """
    stats = {}

    if progress_callback:
        progress_callback(2, 100, "Starting AI rewrite...")

    # ── STAGE 1: Neural Rewrite (core) ───────────────────────────────────────
    def _pass19_cb(curr, tot, msg):
        if progress_callback:
            pct = 5 + int(60 * (curr / max(tot, 1)))
            progress_callback(pct, 100, msg)

    text, n = pass19_structural_smoothing(text, progress_callback=_pass19_cb, register=register)
    stats['pass19_structural_smoothing'] = n

    if progress_callback:
        progress_callback(66, 100, "Removing AI fingerprints...")

    # ── STAGE 2: Light cleanup (invisible, safe) ─────────────────────────────
    text, n = pass1_ai_phrases(text);        stats['pass1_ai_phrases'] = n
    text, n = pass5_contractions(text);      stats['pass5_contractions'] = n
    text, n = pass8_hedging(text);           stats['pass8_hedging'] = n
    text, n = pass2_intensifiers(text);      stats['pass2_intensifiers'] = n

    if progress_callback:
        progress_callback(74, 100, "Breaking statistical patterns...")

    # ── STAGE 3: Structural perturbation (defeats neural classifiers) ────────
    # These passes change the token-level patterns that GPTZero/DeBERTa detect.
    # They run for ALL model types — the neural rewrite alone is NOT enough.

    # Burstiness: split long sentences, merge short ones → rhythm variation
    text, n = pass3_burstiness(text);        stats['pass3_burstiness'] = n

    # Perplexity-guided perturbation: use GPT-2 to find sentences that are
    # "too predictable" (low perplexity = AI-like) and aggressively swap words
    # in those sentences using student-level WordNet synonyms.
    # This is the KEY pass — it directly attacks what GPTZero measures.
    text, n = pass15_perplexity_guided_perturbation(text)
    stats['pass15_synonym_rotation'] = n

    # Passive → active voice: structural transformation
    text, n = pass6_passive_to_active(text); stats['pass6_passive_voice'] = n

    # Punctuation humanization: em-dashes, semicolons disrupt token boundaries
    text, n = pass13_punctuation(text, register=register);      stats['pass13_punctuation'] = n

    # Syntactic fronting: move clauses around to break SVO patterns
    text, n = pass14_syntactic_fronting(text); stats['pass14_syntactic_fronting'] = n

    if progress_callback:
        progress_callback(85, 100, "Adding human signals...")

    # Imperfect discourse: conjunction openers (And, But, Yet) — human-only
    text, n = pass16_imperfect_discourse(text); stats['pass16_imperfect_discourse'] = n

    # Opener diversity: prevent same paragraph opener repeating
    text, n = pass7_opener_diversity(text);  stats['pass7_opener_diversity'] = n

    # Perplexity tension: inject organic idioms to shatter predictability
    text, n = pass18_perplexity_tension(text, register=register); stats['pass18_perplexity_tension'] = n

    # Ghost character cleanup
    text, n = pass17_ghost_characters(text); stats['pass17_ghost_characters'] = n

    # Final cleanup: revert any em-dashes introduced during Stage 3 or neural rewrite
    text, n = pass20_em_dash_reversal(text); stats['pass20_em_dash_reversal'] = n

    if progress_callback:
        progress_callback(95, 100, "Finalizing output...")

    # Zeroed-out passes (still too fingerprint-prone for GPTZero)
    for key in [
        'pass4_discourse_markers', 'pass9_parentheticals',
        'pass10_self_corrections', 'pass11_personal_voice',
        'pass12_qualifiers',
    ]:
        stats[key] = 0

    stats['total_changes'] = sum(stats.values())
    return {'text': text.strip(), 'stats': stats}


# ─────────────────────────────────────────────────────────────────────────────
# DOCX-LEVEL PROCESSING
# ─────────────────────────────────────────────────────────────────────────────

def is_code_paragraph(text: str) -> bool:
    """Detect code paragraphs that should be skipped."""
    s = text.strip()
    return any([
        s.startswith('//'), s.startswith('import '), s.startswith('export '),
        s.startswith('const '), s.startswith('function '), s.startswith('class '),
        s.startswith('{'), s.startswith('}'), s.startswith('return '),
        text.count('{') > 1, text.count('(') > 2,
        '=>' in text, '::' in text, 'await ' in text, 'async ' in text,
    ])


def is_heading_text(text: str) -> bool:
    """Detect short heading-like paragraphs to skip."""
    s = text.strip()
    return len(s.split()) <= 6 and not s.endswith('.')


def humanize_docx(input_path: str, output_path: str, progress_callback=None) -> dict:
    """
    Read a .docx, run 17-pass humanization on all prose paragraphs, save output.
    Returns stats dict.
    """
    from docx import Document

    doc = Document(input_path)
    all_stats = {
        'pass1_ai_phrases': 0, 'pass2_intensifiers': 0,
        'pass3_burstiness': 0, 'pass4_discourse_markers': 0,
        'pass5_contractions': 0, 'pass6_passive_voice': 0,
        'pass7_opener_diversity': 0, 'pass8_hedging': 0,
        'pass9_parentheticals': 0, 'pass10_self_corrections': 0,
        'pass11_personal_voice': 0, 'pass12_qualifiers': 0,
        'pass13_punctuation': 0, 'pass14_syntactic_fronting': 0,
        'pass15_synonym_rotation': 0,
        'pass16_imperfect_discourse': 0,
        'pass17_ghost_characters': 0,
        'pass18_perplexity_tension': 0,
        'pass19_structural_smoothing': 0,
        'pass20_em_dash_reversal': 0,
        'total_changes': 0,
        'paragraphs_processed': 0,
        'paragraphs_skipped': 0,
    }

    def is_valid_para(para_text):
        original = para_text.strip()
        if not original or len(original.split()) < 5:
            return False
        if is_heading_text(original) or is_code_paragraph(original):
            return False
        return True

    valid_paras = []
    for para in doc.paragraphs:
        if is_valid_para(para.text):
            valid_paras.append(para)
            
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if is_valid_para(para.text):
                        valid_paras.append(para)

    total_valid = len(valid_paras)
    
    if progress_callback:
        progress_callback(5, 100, f"Found {total_valid} valid paragraphs. Initializing AI...")

    for i, para in enumerate(valid_paras):
        if progress_callback:
            pct = 10 + int(85 * (i / max(total_valid, 1)))
            progress_callback(pct, 100, f"Processing paragraph {i+1} of {total_valid}...")
            
        original = para.text.strip()
        result = humanize_text(original)
        new_text = result['text']

        if new_text != original:
            if para.runs:
                para.runs[0].text = new_text
                for run in para.runs[1:]:
                    run.text = ''
            else:
                para.text = new_text

            for k, v in result['stats'].items():
                if k in all_stats:
                    all_stats[k] += v
            all_stats['paragraphs_processed'] += 1

    all_stats['paragraphs_skipped'] = len(doc.paragraphs) + sum(len(c.paragraphs) for t in doc.tables for r in t.rows for c in r.cells) - total_valid

    all_stats['total_changes'] = sum(
        v for k, v in all_stats.items() if k.startswith('pass')
    )

    # ── FORENSIC SCRUBBER: METADATA FORGERY ──
    import datetime
    core_props = doc.core_properties
    
    # Randomize revision count to look like a heavily edited human document
    core_props.revision = random.randint(45, 130)
    
    # Backdate the creation time to look like it took days to write
    days_ago = random.randint(3, 10)
    created_time = datetime.datetime.now(datetime.timezone.utc) - datetime.timedelta(days=days_ago, hours=random.randint(1, 12))
    core_props.created = created_time
    core_props.modified = datetime.datetime.now(datetime.timezone.utc)
    
    # Clean up author tags if they look like AI or system tags
    author_name = core_props.author or ""
    if "GPT" in author_name or "AI" in author_name or "Admin" in author_name:
        core_props.author = "User"
        
    core_props.last_modified_by = core_props.author

    doc.save(output_path)
    return all_stats

# Triggering rebuild to force local LoRA loading.
